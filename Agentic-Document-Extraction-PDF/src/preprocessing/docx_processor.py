"""
DOCX file processor.

Converts Microsoft Word documents into page images by rendering
each page to a PNG image for VLM extraction.

Requires: python-docx>=1.1.0
"""

import io
import secrets
import time
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from PIL import Image, ImageDraw, ImageFont

from src.config import get_logger
from src.preprocessing.base_processor import (
    BaseFileProcessor,
    FileProcessingError,
    FileValidationError,
)
from src.preprocessing.pdf_processor import (
    PageImage,
    PDFMetadata,
    ProcessingResult,
)


logger = get_logger(__name__)

# Standard US Letter page at 300 DPI
PAGE_WIDTH_PX = 2550   # 8.5in * 300dpi
PAGE_HEIGHT_PX = 3300  # 11in * 300dpi
MARGIN_PX = 150        # ~0.5in margin
LINE_HEIGHT = 36
FONT_SIZE = 28
MAX_FILE_SIZE_MB = 100


class DocxProcessor(BaseFileProcessor):
    """
    Process DOCX files into PageImage objects.

    Extracts text content from DOCX paragraphs and tables, renders them
    as page images for VLM processing. Layout is approximate — the VLM
    handles understanding the visual structure.
    """

    def __init__(self, dpi: int = 300) -> None:
        self._dpi = dpi

    def validate(self, file_path: Path) -> None:
        """Validate DOCX file."""
        if not file_path.exists():
            raise FileValidationError(f"File not found: {file_path}")

        suffix = file_path.suffix.lower()
        if suffix not in (".docx", ".doc"):
            raise FileValidationError(f"Not a Word document: {suffix}")

        size_mb = file_path.stat().st_size / (1024 * 1024)
        if size_mb > MAX_FILE_SIZE_MB:
            raise FileValidationError(
                f"File size {size_mb:.1f} MB exceeds limit of {MAX_FILE_SIZE_MB} MB"
            )

    def process(self, file_path: Path) -> ProcessingResult:
        """Convert DOCX to page images by rendering text content."""
        start_time = time.monotonic()
        self.validate(file_path)

        try:
            from docx import Document as DocxDocument
        except ImportError as e:
            raise FileProcessingError(
                "python-docx is required for DOCX processing. "
                "Install with: pip install python-docx>=1.1.0"
            ) from e

        file_hash = self._compute_file_hash(file_path)
        processing_id = f"docx_{secrets.token_hex(8)}"
        warnings: list[str] = []

        doc = DocxDocument(str(file_path))

        # Extract all text lines from paragraphs and tables
        text_lines = self._extract_text_lines(doc)

        # Render text lines into page images
        pages = self._render_text_to_pages(text_lines)

        if not pages:
            warnings.append("No content found in DOCX file")

        processing_time_ms = int((time.monotonic() - start_time) * 1000)

        metadata = PDFMetadata(
            file_path=file_path,
            file_name=file_path.name,
            file_size_bytes=file_path.stat().st_size,
            file_hash=file_hash,
            page_count=len(pages),
            title=doc.core_properties.title or None,
            author=doc.core_properties.author or None,
            subject=doc.core_properties.subject or None,
            keywords=doc.core_properties.keywords or None,
            creator="DocxProcessor",
            producer="python-docx",
            creation_date=doc.core_properties.created or datetime.now(UTC),
            modification_date=doc.core_properties.modified,
            pdf_version="N/A",
            is_encrypted=False,
            has_forms=False,
            has_annotations=False,
            processing_id=processing_id,
        )

        logger.info(
            "docx_processing_complete",
            file=file_path.name,
            pages=len(pages),
            time_ms=processing_time_ms,
        )

        return ProcessingResult(
            metadata=metadata,
            pages=pages,
            processing_time_ms=processing_time_ms,
            warnings=warnings,
        )

    def _extract_text_lines(self, doc: Any) -> list[str]:
        """Extract text lines from DOCX paragraphs and tables."""
        lines: list[str] = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Paragraph
                text = element.text or ""
                for child in element.iter():
                    if child.text:
                        text = child.text
                    if child.tail:
                        text += child.tail
                # Use the paragraph's full text via python-docx API

            elif tag == "tbl":
                # Table — render as text grid
                lines.append("")  # blank line before table

        # Fallback: use python-docx paragraph API for reliable text extraction
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
            else:
                lines.append("")  # preserve blank lines for spacing

        # Extract table content
        for table in doc.tables:
            lines.append("")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))
            lines.append("")

        return lines

    def _render_text_to_pages(self, text_lines: list[str]) -> list[PageImage]:
        """Render text lines into page images."""
        pages: list[PageImage] = []
        usable_height = PAGE_HEIGHT_PX - 2 * MARGIN_PX
        lines_per_page = usable_height // LINE_HEIGHT

        # Split lines into page-sized chunks
        page_chunks: list[list[str]] = []
        for i in range(0, max(len(text_lines), 1), lines_per_page):
            chunk = text_lines[i : i + lines_per_page]
            if chunk:
                page_chunks.append(chunk)

        if not page_chunks:
            page_chunks = [[""]]

        for page_num, chunk in enumerate(page_chunks, start=1):
            page_img = self._render_page(chunk, page_num)
            pages.append(page_img)

        return pages

    def _render_page(self, lines: list[str], page_number: int) -> PageImage:
        """Render a single page of text as an image."""
        img = Image.new("RGB", (PAGE_WIDTH_PX, PAGE_HEIGHT_PX), "white")
        draw = ImageDraw.Draw(img)

        try:
            font = ImageFont.truetype("arial.ttf", FONT_SIZE)
        except OSError:
            font = ImageFont.load_default()

        y = MARGIN_PX
        text_content_parts: list[str] = []

        for line in lines:
            draw.text((MARGIN_PX, y), line, fill="black", font=font)
            text_content_parts.append(line)
            y += LINE_HEIGHT

        # Convert to PNG bytes
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        png_bytes = buf.getvalue()

        orientation = self._detect_orientation(PAGE_WIDTH_PX, PAGE_HEIGHT_PX)

        return PageImage(
            page_number=page_number,
            image_bytes=png_bytes,
            width=PAGE_WIDTH_PX,
            height=PAGE_HEIGHT_PX,
            dpi=self._dpi,
            orientation=orientation,
            original_width_pts=612.0,   # US Letter: 8.5 * 72
            original_height_pts=792.0,  # US Letter: 11 * 72
            has_text=True,
            has_images=False,
            rotation=0,
            text_content="\n".join(text_content_parts),
        )
